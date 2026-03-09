
import os
from pdf2docx import Converter
import fitz  # 处理pdf的工具 pymupdf
import re
from docx import Document


# 检验pdf的可用性
# def 函数名(参数名: 类型) -> 返回值类型：
def validate_pdf_file(file_path: str) -> bool:
    """
        Args: file_path: 检查pdf文件路径
        Returns:可用true 不可用false
    """

    # 判断文件是否存在
    if not os.path.exists(file_path):
        print(f"【错误】文件不存在：{file_path}")
        return False

    # 判断文件是否为pdf lower将路径转为小写 endswith 判断后缀
    if not file_path.lower().endswith(".pdf"):
        print(f"【错误】不是pdf文件：{file_path}")
        return False

    # 判断文件能否打开
    try:
        fitz.open(file_path)
        return True
    except Exception as e:
        print(f"【错误】文件无法打开：{file_path}")
        return False


# 将pdf转为docx文件
def pdf_2_docx(pdf_file: str,docx_file:str) -> str:
    if not os.path.exists(pdf_file):
        print(f"【错误】pdf文件不存在")
        return " "
    cv = Converter(pdf_file)
    cv.convert(docx_file)
    cv.close()
    if os.path.exists(docx_file):
        print(f"【成功】pdf转word成功/n")
    return docx_file

#对docx文件进行切分
'''
        1.按段进行切分
        2.滤除噪声：纯数字 or 注解
        3，分页导致断段 拼接
'''

# 3. 噪声匹配规则：针对法律PDF的脚注/页码/超链接
FOOTNOTE_PATTERN = re.compile(r'^\d+\s+可參閱網址:')  # 匹配"1 可參閱網址:"类脚注
PAGE_NUM_PATTERN = re.compile(r'^\d+$')
# 1. 句末标点：判断句子是否结束的核心标记
SENTENCE_END_MARKS = {"。", "！", "？", "；", ".", "!", "?", ";"}
def process_docx(docx_file: str) -> list[str]:
    #检验文件是否存在
    try:
        doc = Document(docx_file)
        print(f"[成功】docx文件成功读取")
    except Exception as e:
        print(f"docx文件不存在，{e}")

    #读取所有段落
    raw_paras=[]
    for para in doc.paragraphs:
        para_text = para.text.strip() #清楚段落中的多余空格符
        if para_text:
            raw_paras.append(para_text)
    print(f"已读取原始段落，共“{len(raw_paras)}段")

    # 滤除噪声
    cleaned_paras = []
    for para in raw_paras:
        # 过滤纯数字
        if PAGE_NUM_PATTERN.match(para):
            continue
        # 过滤脚注
        if FOOTNOTE_PATTERN.match(para):
            continue
        else:
            cleaned_paras.append(para)
    print(f"🧹 过滤噪声完成 → 共有{len(cleaned_paras)}段")

    # 段落拼接
    if not cleaned_paras:
        return []

    clean_paras = []
    current_sentence = cleaned_paras[0] #初始化第一段

    #遍历段落 判断是否结束
    for para in cleaned_paras[1:]:
        #句子未结束 判断最后一个字符
        if current_sentence  and current_sentence[-1] not in SENTENCE_END_MARKS:
            current_sentence += para
        else:
            clean_paras.append(current_sentence)
            current_sentence = para
    #将最后一句存入
    if current_sentence:
        clean_paras.append(current_sentence)

    print(f"合并断句完成->原始{len(cleaned_paras)}段 → 合并后{len(clean_paras)}段")


    return clean_paras

# 分块 按句拆分
def split_text_to_chunks(clean_paras:list[str],max_length = 300) -> list[dict]:
    chunks = []
    for para in clean_paras:

        #清理空格
        para_clean = para.strip()
        if not para_clean:
            continue

        # 段落长度
        if len(para_clean) > max_length:
            sentences =  re.split(r'([。！？])', para_clean) # 正则化拆分
            # 拆分结果示例：["貼科技發展和掌握潛在的科技罪行趨勢", "。", "下一步要加強監管", "！"]
            #将标点与句子拼接
            complete_sentences = []
            for i in range(0,len(sentences)-1,2):
                if sentences[i]:
                    complete_sentences.append(sentences[i]+sentences[i+1])
           #将完整句子添加到分块中
            for sentence in complete_sentences:
                chunks.append({
                    "text": sentence.strip(),
                    "format_info":{"page":None}
                })
        else:
            chunks.append({
                    "text": para_clean,
                    "format_info":{"page":None}
                })
    print(f"✅ 分块完成 → 共{len(chunks)}个翻译块（每块最多{max_length}字符）")
    return chunks


#翻译结果写入docx
def write_translation_docx(translated_chunks:list[dict],output_path: str) -> None:
    doc = Document()
    for chunk in translated_chunks:
        doc.add_paragraph(chunk["text"])
    doc.save(output_path)
    print(f"✅ 翻译结果已写入 → {output_path}")



# 测试代码
if __name__ == "__main__":
    test_pdf = "../document/检控科.pdf"
    doc_file = "../test/pdf_2_doc.docx"
    output_path="../test/测试结果.docx"

    print("\n【步骤1】PDF转DOCX")
    docx_result  =pdf_2_docx(test_pdf,doc_file)
    if not docx_result:
        print("❌ 测试失败：PDF转DOCX失败")

    print("\n【步骤2】DOCX去噪声+合并断句")
    clean_paras = process_docx(doc_file)
    if not clean_paras:
        print("❌ 测试失败：DOCX处理无有效段落")
    print("📝 干净段落详情：")
    for i, para in enumerate(clean_paras):
        print(f"  段落{i + 1}：{para}（长度：{len(para)}）")

    print("\n【步骤3】文本分块")
    chunks = split_text_to_chunks(clean_paras, max_length = 300)
    if not chunks:
        print("❌ 测试失败：分块无有效内容")
    print("📝 分块详情：")
    for i, chunk in enumerate(chunks):
        print(f"  块{i + 1}：{chunk['text']}（长度：{len(chunk['text'])}）")


    print("\n【步骤4】模拟翻译结果写入DOCX")
    mock_translated = []
    for i, chunk in enumerate(chunks):
        mock_trans = f"[翻译测试] {chunk['text']}"
        mock_translated.append({"text": mock_trans, "format_info": chunk["format_info"]})
    # 写入DOCX
    write_translation_docx(mock_translated, output_path)
